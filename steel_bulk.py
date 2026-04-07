# -*- coding: utf-8 -*-
"""
Abter Steel — 群开系统
目标：拉美（巴西+墨西哥+智利）
输出：按行业分类的公司名单 + 每个行业一封通用开发信
成本极低：搜索用Serper，写信用Claude，不做个性化
"""
import os
import json
import re
import requests
from crewai import Agent, Task, Crew, LLM
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt
from datetime import datetime
load_dotenv()

my_llm = LLM(
    model="openai/claude-sonnet-4-6",
    api_key=os.environ["CLAUDE_API_KEY"],
    base_url=os.environ["CLAUDE_API_URL"],
    temperature=0.4
)

# ══════════════════════════════════════════
# 配置：目标国家 + 行业
# ══════════════════════════════════════════
TARGET_COUNTRIES = ["Brazil", "Mexico", "Chile"]

INDUSTRIES = {
    "oil_gas":      {"name": "石油天然气",     "en": "Oil & Gas",          "es": "Petróleo y Gas"},
    "mining":       {"name": "矿业",           "en": "Mining",             "es": "Minería"},
    "construction": {"name": "建筑/EPC承包商", "en": "Construction & EPC", "es": "Construcción y EPC"},
    "petrochemical":{"name": "石化化工",       "en": "Petrochemical",      "es": "Petroquímica"},
    "trader":       {"name": "钢管贸易商",     "en": "Steel Pipe Trader",  "es": "Comerciante de Tubos"},
}

# ══════════════════════════════════════════
# 工具函数
# ══════════════════════════════════════════
def serper_search(query, retries=3):
    for attempt in range(retries):
        try:
            resp = requests.post(
                "https://google.serper.dev/search",
                headers={"X-API-KEY": os.environ["SERPER_API_KEY"]},
                json={"q": query, "num": 8},
                timeout=15,
            )
            results = resp.json().get("organic", [])
            return "\n".join([
                f"- {r.get('title','')}: {r.get('link','')} | {r.get('snippet','')[:80]}"
                for r in results
            ])
        except Exception as e:
            if attempt < retries - 1:
                print(f"    ⚠️ 网络中断，3秒后重试({attempt+1}/{retries})...")
                import time
                time.sleep(3)
            else:
                print(f"    ❌ 搜索失败，跳过：{query[:30]}")
                return ""
    return ""
    
def parse_json(text):
    try:
        match = re.search(r'\[.*\]', text, re.DOTALL)
        return json.loads(match.group()) if match else []
    except:
        return []

# ══════════════════════════════════════════
# Agent 定义
# ══════════════════════════════════════════
extractor = Agent(
    role="公司信息提取专家",
    goal="从搜索结果中提取公司名称和官网地址",
    backstory="你擅长从搜索结果中快速识别真实公司并提取其网址，过滤掉新闻和无关内容。",
    llm=my_llm, verbose=False, allow_delegation=False,
)

email_writer = Agent(
    role="西班牙语外贸开发信专家",
    goal="为中国钢管出口商撰写面向拉美市场的专业西班牙语开发信",
    backstory=(
        "你是Abter Steel拉美市场开发信专家，熟悉巴西、墨西哥、智利的商业文化。"
        "你的信件专业、简洁，针对不同行业突出不同的产品优势。"
        "你用西班牙语写作，语气正式但不失亲切。"
    ),
    llm=my_llm, verbose=False, allow_delegation=False,
)

# ══════════════════════════════════════════
# 主流程
# ══════════════════════════════════════════
print("=" * 60)
print("  Abter Steel — 拉美群开系统启动")
print(f"  目标国家：{' + '.join(TARGET_COUNTRIES)}")
print(f"  目标行业：{len(INDUSTRIES)} 个")
print("=" * 60)

all_companies = {}   # {行业: [{country, company, website}]}
all_emails = {}      # {行业: {country: 开发信}}

# ══════════════════════════════════════════
# 第一阶段：按国家+行业搜索公司
# ══════════════════════════════════════════
print("\n📡 第一阶段：搜索各国各行业公司名单...\n")

for industry_key, industry_info in INDUSTRIES.items():
    all_companies[industry_key] = []
    print(f"  行业：{industry_info['name']}")

    for country in TARGET_COUNTRIES:
        # 构造搜索词（英语+西班牙语双语搜索效果更好）
        query = f"{country} {industry_info['en']} company steel pipe supplier list site:.com OR site:.{country[:2].lower()}"
        print(f"    搜索 {country}...")
        search_result = serper_search(query)

        # 用Agent提取公司名和网址
        extract_task = Task(
            description=(
                f"从以下搜索结果中提取{country}{industry_info['name']}行业的真实公司：\n\n"
                f"{search_result}\n\n"
                f"要求：\n"
                f"1. 只提取真实存在的公司，过滤新闻/博客/目录网站\n"
                f"2. 每个公司要有官网地址\n"
                f"3. 提取3-5家即可\n"
                f"严格JSON输出：\n"
                f'[{{"company":"公司名","website":"https://...","country":"{country}","industry":"{industry_info["en"]}"}}]'
            ),
            expected_output="JSON数组，3-5家公司",
            agent=extractor,
        )
        companies = parse_json(
            str(Crew(agents=[extractor], tasks=[extract_task], verbose=False).kickoff())
        )
        all_companies[industry_key].extend(companies)
        print(f"    ✅ {country}：找到 {len(companies)} 家")

total = sum(len(v) for v in all_companies.values())
print(f"\n  ✅ 第一阶段完成，共找到 {total} 家公司")

# ══════════════════════════════════════════
# 第二阶段：每个行业生成一封通用西班牙语开发信
# ══════════════════════════════════════════
print("\n✉️  第二阶段：生成各行业通用开发信（西班牙语）...\n")

industry_contexts = {
    "oil_gas":       "石油天然气开采和管道运输，需要API 5L管线管、油套管、结构管",
    "mining":        "矿山开采和矿浆输送，需要耐磨耐压的无缝钢管和结构钢管",
    "construction":  "大型基建项目承包，需要结构空心型钢和施工用钢管",
    "petrochemical": "炼油和化工装置，需要高压耐腐蚀的合金管和不锈钢管",
    "trader":        "钢管贸易分销，需要多规格库存管材，价格和交货期优先",
}

for industry_key, industry_info in INDUSTRIES.items():
    context = industry_contexts.get(industry_key, "")
    companies_count = len(all_companies.get(industry_key, []))
    print(f"  生成 [{industry_info['name']}] 行业开发信...")

    email_task = Task(
        description=(
            f"为拉美{industry_info['name']}行业客户写一封西班牙语开发信（150-200词）\n\n"
            f"行业背景：{context}\n"
            f"目标国家：{', '.join(TARGET_COUNTRIES)}\n\n"
            f"发件人：Abter Steel，中国河北沧州\n"
            f"核心优势：年产能12万吨，API 5L/ISO 9001认证，拉美45天交货\n"
            f"联系方式：sales@abter-steel.com | WhatsApp: +86-139-XXXX-XXXX\n\n"
            f"要求：\n"
            f"1. 开头直接点明了解该行业对钢管的具体需求\n"
            f"2. 突出对该行业最相关的产品和认证\n"
            f"3. 结尾提供免费报价，附WhatsApp联系方式\n"
            f"4. 用正式西班牙语（usted格式）\n"
            f"5. 不要填具体公司名，用'Estimado responsable de compras'开头\n\n"
            f"格式：\nAsunto: [...]\nEstimado responsable de compras,\n[正文]\nAtentamente,\n[签名]"
        ),
        expected_output="完整西班牙语开发信，含Asunto行和正文",
        agent=email_writer,
    )
    email_text = str(Crew(agents=[email_writer], tasks=[email_task], verbose=False).kickoff())
    all_emails[industry_key] = email_text
    print(f"  ✅ [{industry_info['name']}] 完成")

# ══════════════════════════════════════════
# 第三阶段：导出 Word 文档
# ══════════════════════════════════════════
print("\n📄 第三阶段：导出Word文档...")

doc = Document()

# 封面
doc.add_heading("Abter Steel — 拉美市场群开报告", 0)
doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y年%m月%d日 %H:%M')}")
doc.add_paragraph(f"目标国家：{' | '.join(TARGET_COUNTRIES)}")
doc.add_paragraph(f"覆盖行业：{len(INDUSTRIES)} 个  |  公司总数：{total} 家")
doc.add_page_break()

# 按行业输出：公司名单 + 通用开发信
for industry_key, industry_info in INDUSTRIES.items():
    companies = all_companies.get(industry_key, [])
    email = all_emails.get(industry_key, "")

    doc.add_heading(f"{industry_info['name']} / {industry_info['es']}", level=1)

    # 公司名单表格
    doc.add_heading("客户名单（人工补充联系人）", level=2)
    if companies:
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        for cell, text in zip(table.rows[0].cells, ['公司名', '国家', '官网']):
            cell.text = text
            cell.paragraphs[0].runs[0].bold = True

        for c in companies:
            row = table.add_row().cells
            row[0].text = c.get('company', '')
            row[1].text = c.get('country', '')
            row[2].text = c.get('website', '')
    else:
        doc.add_paragraph("（未找到公司数据）")

    doc.add_paragraph("")

    # 通用开发信
    doc.add_heading("通用开发信（西班牙语）", level=2)
    for line in email.split('\n'):
        line = line.strip().replace('**', '')
        if line:
            p = doc.add_paragraph(line)
            p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()

filename = f"abter_bulk_latam_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
doc.save(filename)

print(f"\n{'=' * 60}")
print(f"  ✅ 群开系统完成！")
print(f"  📁 报告文件：{filename}")
print(f"  🌎 覆盖国家：{len(TARGET_COUNTRIES)} 个")
print(f"  🏭 覆盖行业：{len(INDUSTRIES)} 个")
print(f"  🏢 公司总数：{total} 家")
print(f"  ✉️  行业开发信：{len(all_emails)} 封")
print(f"{'=' * 60}\n")