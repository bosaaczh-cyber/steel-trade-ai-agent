# -*- coding: utf-8 -*-
"""
Abter Steel — 全自动外贸流程
运行一次：搜索客户 → 评分 → 对A/B级生成开发信 → 导出完整报告
"""
import os
import json
import re
import requests
from crewai import Agent, Task, Crew, LLM
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, RGBColor
from datetime import datetime
load_dotenv()

my_llm = LLM(
    model="openai/claude-sonnet-4-6",
    api_key=os.environ["CLAUDE_API_KEY"],
    base_url=os.environ["CLAUDE_API_URL"],
    temperature=0.4
)

# ══════════════════════════════════════════
# 工具函数
# ══════════════════════════════════════════
def serper_search(query):
    resp = requests.post(
        "https://google.serper.dev/search",
        headers={"X-API-KEY": os.environ["SERPER_API_KEY"]},
        json={"q": query, "num": 5}
    )
    results = resp.json().get("organic", [])
    return "\n".join([f"- {r['title']}: {r.get('snippet','')}" for r in results])

def parse_json(text):
    try:
        match = re.search(r'\[.*\]', text, re.DOTALL)
        return json.loads(match.group()) if match else []
    except:
        return []

def print_banner(title):
    print(f"\n{'=' * 55}")
    print(f"  {title}")
    print(f"{'=' * 55}\n")

# ══════════════════════════════════════════
# Agent 定义（一次性创建，全程复用）
# ══════════════════════════════════════════
researcher = Agent(
    role="石油化工行业外贸客户研究员",
    goal="从搜索结果中提取结构化客户线索",
    backstory="你是Abter Steel的资深外贸研究员，专门从公开信息中提取真实采购线索。",
    llm=my_llm, verbose=False, allow_delegation=False,
)

scorer = Agent(
    role="外贸线索质量评分专家",
    goal="对钢管外贸线索进行客观评分，输出优先级排序",
    backstory="你是Abter Steel外贸总监，15年钢管出口经验，评分客观严格。",
    llm=my_llm, verbose=False, allow_delegation=False,
)

writer = Agent(
    role="专业外贸开发信撰写专家",
    goal="为高价值客户撰写高转化率英语开发信",
    backstory=(
        "你是Abter Steel资深开发信专家，10年B2B钢铁出口经验。"
        "开头直接引用客户具体项目，用数字说话，绝不用废话开头。"
    ),
    llm=my_llm, verbose=False, allow_delegation=False,
)

# ══════════════════════════════════════════
# 第一阶段：搜索找客户
# ══════════════════════════════════════════
print_banner("Abter Steel — 全自动外贸流程启动")

queries = [
    "Saudi Aramco steel pipe procurement manager LinkedIn",
    "Middle East oil company steel pipe supplier China 2024",
    "ADNOC SABIC steel pipe procurement contact",
]

print("📡 第一阶段：搜索客户线索...")
all_search_results = ""
for i, q in enumerate(queries, 1):
    print(f"  搜索 {i}/3: {q[:45]}...")
    result = serper_search(q)
    all_search_results += f"\n【搜索{i}】{q}\n{result}\n"

research_task = Task(
    description=(
        f"搜索结果如下：\n{all_search_results}\n\n"
        "整理出5家公司，严格按JSON输出，不要其他文字：\n"
        '[{"company":"公司名","country":"国家","contact":"联系人或职位",'
        '"title":"职位","need":"钢管需求描述"}]'
    ),
    expected_output="JSON数组，5家公司",
    agent=researcher,
)
leads = parse_json(str(Crew(agents=[researcher], tasks=[research_task], verbose=False).kickoff()))

# 解析失败用默认数据
if not leads:
    leads = [
        {"company":"Saudi Aramco","country":"Saudi Arabia","contact":"Mohammad Riaz","title":"Procurement Manager","need":"RCTF项目结构钢及机械固定设备采购"},
        {"company":"ADNOC","country":"UAE","contact":"Procurement Manager","title":"Procurement Manager","need":"1600km管道网络精炼产品管材"},
        {"company":"East Pipes","country":"Saudi Arabia","contact":"未知","title":"未知","need":"5000万美元钢管供应协议"},
        {"company":"SABIC","country":"Saudi Arabia","contact":"未知","title":"未知","need":"石化工厂设备用管材"},
        {"company":"Kuwait Oil Company","country":"Kuwait","contact":"未知","title":"未知","need":"油田开发钢管需求"},
    ]

print(f"  ✅ 找到 {len(leads)} 家目标客户")

# ══════════════════════════════════════════
# 第二阶段：客户评分
# ══════════════════════════════════════════
print("\n📊 第二阶段：客户评分中...")

score_task = Task(
    description=(
        f"对以下线索评分：\n{json.dumps(leads, ensure_ascii=False)}\n\n"
        "评分标准：公司规模(30分)+需求明确度(25分)+联系人质量(25分)+市场潜力(20分)\n"
        "A级=80+分，B级=60-79分，C级=60分以下\n"
        "严格JSON输出：\n"
        '[{"company":"名","total_score":85,"grade":"A",'
        '"score_breakdown":{"company_size":28,"need_clarity":23,"contact_quality":20,"market_potential":14},'
        '"priority":"立即跟进","reason":"原因","suggested_action":"行动建议"}]'
    ),
    expected_output="JSON评分数组",
    agent=scorer,
)
scored_leads = parse_json(str(Crew(agents=[scorer], tasks=[score_task], verbose=False).kickoff()))
scored_leads.sort(key=lambda x: x.get('total_score', 0), reverse=True)

# 打印评分结果
grade_icons = {"A": "⭐⭐⭐", "B": "⭐⭐", "C": "⭐"}
print(f"\n  {'排名':<4} {'评级':<10} {'分数':<6} {'公司'}")
print(f"  {'-'*45}")
for i, lead in enumerate(scored_leads, 1):
    g = lead.get('grade','C')
    print(f"  {i:<4} {g}级{grade_icons.get(g,''):<8} {lead.get('total_score',0):<6} {lead.get('company','')}")

a_leads = [l for l in scored_leads if l.get('grade') == 'A']
b_leads = [l for l in scored_leads if l.get('grade') == 'B']
c_leads = [l for l in scored_leads if l.get('grade') == 'C']
print(f"\n  A级：{len(a_leads)}家 | B级：{len(b_leads)}家 | C级：{len(c_leads)}家（C级跳过开发信）")

# ══════════════════════════════════════════
# 第三阶段：只对A/B级生成开发信
# ══════════════════════════════════════════
priority_leads = [l for l in scored_leads if l.get('grade') in ['A', 'B']]
print(f"\n✉️  第三阶段：为 {len(priority_leads)} 家A/B级客户生成开发信...")

# 把评分数据和原始leads合并
leads_dict = {l['company']: l for l in leads}
email_results = []

for i, scored in enumerate(priority_leads, 1):
    company = scored.get('company', '')
    original = leads_dict.get(company, scored)
    contact = original.get('contact', 'Procurement Manager')
    title = original.get('title', 'Procurement Manager')
    country = original.get('country', '')
    need = original.get('need', '')
    grade = scored.get('grade', 'B')
    action = scored.get('suggested_action', '')

    print(f"  生成 {i}/{len(priority_leads)}: {company} ({grade}级)...")

    email_task = Task(
        description=(
            f"为以下{grade}级客户写英语开发信（150-200词）：\n\n"
            f"收件人：{contact}，{title}\n"
            f"公司：{company}，{country}\n"
            f"需求背景：{need}\n"
            f"我方建议行动：{action}\n\n"
            f"发件人：Abter Steel，河北沧州\n"
            f"核心优势：年产能12万吨，API 5L/ISO 9001认证，中东30天交货\n"
            f"联系方式：sales@abter-steel.com | WhatsApp: +86-139-XXXX-XXXX\n\n"
            f"格式：\nSubject: [...]\nDear {contact},\n[正文]\nBest regards,\n[签名]"
        ),
        expected_output="完整英语开发信",
        agent=writer,
    )
    email_text = str(Crew(agents=[writer], tasks=[email_task], verbose=False).kickoff())
    email_results.append({
        "scored": scored,
        "original": original,
        "email": email_text
    })
    print(f"  ✅ {company} 完成")

# ══════════════════════════════════════════
# 第四阶段：导出完整 Word 报告
# ══════════════════════════════════════════
print("\n📄 第四阶段：导出完整报告...")

doc = Document()

# 封面
doc.add_heading("Abter Steel — 全自动外贸报告", 0)
doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y年%m月%d日 %H:%M')}")
doc.add_paragraph(f"搜索客户：{len(leads)} 家  |  A级：{len(a_leads)} 家  |  B级：{len(b_leads)} 家  |  C级：{len(c_leads)} 家")
doc.add_paragraph(f"已生成开发信：{len(email_results)} 封（仅A/B级）")
doc.add_page_break()

# 评分汇总表
doc.add_heading("第一部分：客户评分汇总", level=1)
table = doc.add_table(rows=1, cols=5)
table.style = 'Table Grid'
for cell, text in zip(table.rows[0].cells, ['排名','评级','总分','公司','建议行动']):
    cell.text = text
    cell.paragraphs[0].runs[0].bold = True

for i, lead in enumerate(scored_leads, 1):
    row = table.add_row().cells
    row[0].text = str(i)
    row[1].text = f"{lead.get('grade','C')}级"
    row[2].text = str(lead.get('total_score', 0))
    row[3].text = lead.get('company', '')
    row[4].text = lead.get('suggested_action', '')

doc.add_page_break()

# 开发信正文
doc.add_heading("第二部分：A/B级客户开发信", level=1)
for item in email_results:
    scored = item['scored']
    original = item['original']
    grade = scored.get('grade', 'B')

    doc.add_heading(
        f"{grade}级 | {scored.get('company','')} ({scored.get('total_score',0)}分)",
        level=2
    )
    doc.add_paragraph(f"联系人：{original.get('contact','')}  |  职位：{original.get('title','')}  |  国家：{original.get('country','')}")
    doc.add_paragraph(f"评级原因：{scored.get('reason','')}")
    doc.add_paragraph("")
    doc.add_heading("开发信正文", level=3)

    for line in item['email'].split('\n'):
        line = line.strip().replace('**', '')
        if line:
            p = doc.add_paragraph(line)
            p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()

# C级客户存档
if c_leads:
    doc.add_heading("第三部分：C级客户（暂缓跟进）", level=1)
    for lead in c_leads:
        doc.add_paragraph(
            f"• {lead.get('company','')}（{lead.get('total_score',0)}分）— {lead.get('reason','')}",
        )

filename = f"abter_steel_full_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
doc.save(filename)

# ══ 自动挖掘邮箱 ══
from steel_email_finder import enrich_leads
email_results = enrich_leads(email_results)

# ══ 导出发送CSV ══
import csv
from datetime import datetime

csv_filename = f"abter_steel_full_{datetime.now().strftime('%Y%m%d_%H%M')}.csv"
with open(csv_filename, "w", newline="", encoding="utf-8-sig") as f:
    writer = csv.writer(f)
    writer.writerow(["收件人邮箱", "联系人姓名", "公司", "邮件主题", "邮件正文"])
    for item in email_results:
        original = item["original"]
        email_text = item["email"]
        # 提取Subject和正文
        subject, body = "", []
        for line in email_text.split("\n"):
            line = line.strip().replace("**", "")
            if line.lower().startswith("subject:"):
                subject = line.split(":", 1)[-1].strip()
            elif line:
                body.append(line)
        # 只写有邮箱的
        if original.get("email") and "@" in original.get("email", ""):
            writer.writerow([
                original.get("email", ""),
                original.get("contact", ""),
                original.get("company", ""),
                subject,
                "\n".join(body),
            ])

print(f"  📤 发送CSV已导出 → {csv_filename}")

print(f"\n{'=' * 55}")
print(f"  ✅ 全流程完成！")
print(f"  📁 完整报告：{filename}")
print(f"  👥 搜索客户：{len(leads)} 家")
print(f"  📊 A/B级客户：{len(priority_leads)} 家")
print(f"  ✉️  生成开发信：{len(email_results)} 封")
print(f"  ⏭️  C级跳过：{len(c_leads)} 家")
print(f"{'=' * 55}\n")